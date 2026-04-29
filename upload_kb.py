"""
upload_kb.py — Runtime document uploads into a dedicated ChromaDB collection.

Uploaded documents are stored in the "user_uploads" collection, separate from
the regulatory "framl_kb" collection, so the policy KB is never polluted.
The policy agent merges both collections at retrieval time.
"""

import base64
import io
import os
import sys

import chromadb
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction

_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

from config import CHROMA_PATH

UPLOAD_COLLECTION = "user_uploads"
CHUNK_SIZE    = 150
CHUNK_OVERLAP = 20

# Module-level singletons — created on first use
_ef         = None
_chroma     = None
_collection = None


def _get_collection():
    global _ef, _chroma, _collection
    if _collection is not None:
        return _collection
    if _ef is None:
        _ef = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
    if _chroma is None:
        _chroma = chromadb.PersistentClient(path=CHROMA_PATH)
    try:
        _collection = _chroma.get_collection(UPLOAD_COLLECTION, embedding_function=_ef)
    except Exception:
        _collection = _chroma.create_collection(UPLOAD_COLLECTION, embedding_function=_ef)
    return _collection


def _chunk_text(text: str) -> list:
    words = text.split()
    chunks = []
    for i in range(0, len(words), CHUNK_SIZE - CHUNK_OVERLAP):
        chunk = " ".join(words[i : i + CHUNK_SIZE])
        if chunk.strip():
            chunks.append(chunk)
    return chunks


def _extract_text(data_url: str, file_name: str, file_type: str) -> str:
    """Decode a base64 data-URL and return plain text content."""
    # Strip "data:<mime>;base64," header
    if "," in data_url:
        _, b64 = data_url.split(",", 1)
    else:
        b64 = data_url
    raw = base64.b64decode(b64)

    ext = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""
    ft  = (file_type or "").lower()

    if ext == "pdf" or "pdf" in ft:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(raw))
        pages  = [p.extract_text() or "" for p in reader.pages]
        return " ".join(pages)

    if ext in ("docx", "doc") or "word" in ft or "openxmlformats" in ft:
        import docx as _docx
        doc   = _docx.Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return " ".join(parts)

    # Fallback — plain text / markdown / csv
    return raw.decode("utf-8", errors="replace")


def ingest_upload(data_url: str, file_name: str, file_type: str) -> int:
    """
    Extract text from the uploaded file, chunk it, and store in user_uploads.
    Re-uploading the same filename replaces old chunks.
    Returns the number of chunks stored.
    Raises ValueError if no text can be extracted.
    """
    text = _extract_text(data_url, file_name, file_type)
    if not text.strip():
        raise ValueError(f"No readable text extracted from '{file_name}'.")

    chunks = _chunk_text(text)
    if not chunks:
        raise ValueError(f"Document '{file_name}' produced no usable text chunks.")

    coll = _get_collection()

    # Remove previous version of this file to avoid duplicates on re-upload
    try:
        existing = coll.get(where={"source": file_name})
        if existing.get("ids"):
            coll.delete(ids=existing["ids"])
            print(f"[upload_kb] removed {len(existing['ids'])} old chunks for '{file_name}'")
    except Exception:
        pass

    ids       = [f"upload_{file_name}_{i}" for i in range(len(chunks))]
    metadatas = [{"source": file_name, "type": "upload"} for _ in chunks]

    BATCH = 100
    for i in range(0, len(chunks), BATCH):
        coll.add(
            documents=chunks[i : i + BATCH],
            ids=ids[i : i + BATCH],
            metadatas=metadatas[i : i + BATCH],
        )

    print(f"[upload_kb] ingested '{file_name}' → {len(chunks)} chunks")
    return len(chunks)


def retrieve(query: str, top_k: int = 5) -> tuple:
    """
    Query the user_uploads collection.
    Returns (context_text, sources_list).
    """
    coll = _get_collection()
    count = coll.count()
    if count == 0:
        return "", []
    try:
        results = coll.query(
            query_texts=[query],
            n_results=min(top_k, count),
        )
        docs    = results["documents"][0]
        metas   = results["metadatas"][0]
        sources = list(dict.fromkeys(m["source"] for m in metas))
        sections = [f"[Uploaded document: {m['source']}]\n{d}" for d, m in zip(docs, metas)]
        return "\n\n---\n\n".join(sections), sources
    except Exception as e:
        print(f"[upload_kb] retrieval error: {e}")
        return "", []


def list_uploads() -> list:
    """Return unique filenames currently stored in the user_uploads collection."""
    try:
        coll = _get_collection()
        if coll.count() == 0:
            return []
        result = coll.get()
        seen, names = set(), []
        for m in result.get("metadatas", []):
            src = m.get("source", "")
            if src and src not in seen:
                seen.add(src)
                names.append(src)
        return names
    except Exception:
        return []
